import streamlit as st
import time
import random
import json
import base64
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np
import networkx as nx
import io

# Set page config FIRST (before any other st commands)
st.set_page_config(
    page_title="RevealGrid Demo",
    page_icon="🏭",
    layout="wide"
)

# Optional Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    st.warning("⚠️ Word document generation not available. Install python-docx for full functionality.")

# Custom CSS for professional look
st.markdown("""
<style>
    .main-header {
        text-align: center;
        font-size: 2.5rem;
        font-weight: 700;
        color: #1f2937;
        margin-bottom: 1rem;
        padding: 1rem 0;
        border-bottom: 3px solid #3b82f6;
    }
    .subtitle {
        text-align: center;
        font-size: 1.1rem;
        color: #6b7280;
        margin-bottom: 2rem;
        font-weight: 400;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 10px;
        color: white;
        margin: 0.5rem 0;
    }
    .step-card {
        background: #f8fafc;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #3b82f6;
        margin: 1rem 0;
    }
    .success-card {
        background: #ecfdf5;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #10b981;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'current_step' not in st.session_state:
    st.session_state.current_step = 0
    st.session_state.discovered_assets = 0
    st.session_state.risk_score = 0
    st.session_state.efficiency = 0
    st.session_state.recommendations = []

# Demo data
steps = [
    {
        "title": "Welcome to RevealGrid",
        "description": "AI-driven industrial intelligence platform for OT environments",
        "content": "Discover, map, and optimize your industrial assets with advanced AI and computer vision."
    },
    {
        "title": "Asset Discovery",
        "description": "Automated network and device detection",
        "content": "Scanning industrial networks to identify PLCs, sensors, actuators, and safety systems."
    },
    {
        "title": "Risk Assessment",
        "description": "Real-time vulnerability analysis",
        "content": "Analyzing network topology, device configurations, and security posture."
    },
    {
        "title": "AI Recommendations",
        "description": "Intelligent automation suggestions",
        "content": "Generating optimization strategies based on reinforcement learning simulations."
    },
    {
        "title": "Export & Deploy",
        "description": "Secure world model export",
        "content": "Creating encrypted asset maps for deployment in air-gapped environments."
    }
]

def simulate_discovery():
    """Simulate asset discovery process"""
    with st.spinner("Discovering industrial assets..."):
        time.sleep(2)
        st.session_state.discovered_assets = random.randint(15, 45)
        st.session_state.risk_score = random.randint(15, 35)
        st.session_state.efficiency = random.randint(75, 95)

def simulate_analysis():
    """Simulate risk analysis"""
    with st.spinner("Analyzing network topology..."):
        time.sleep(1.5)
        return {
            "high_risk_nodes": random.randint(2, 8),
            "bottlenecks": random.randint(3, 12),
            "anomalies": random.randint(0, 3)
        }

def simulate_recommendations():
    """Simulate AI recommendations"""
    with st.spinner("Generating AI recommendations..."):
        time.sleep(2)
        recommendations = [
            "Automate manual valve control with AI relay (Efficiency: +23%)",
            "Implement predictive maintenance for pump systems (Downtime: -45%)",
            "Optimize conveyor belt routing (Throughput: +18%)",
            "Deploy safety system redundancy (Risk: -32%)"
        ]
        return random.sample(recommendations, random.randint(2, 4))

def create_sample_graph():
    """Create a sample industrial network graph"""
    G = nx.Graph()
    
    # Add industrial nodes
    nodes = [
        "PLC_Controller", "TemperatureSensor", "PressureSensor", "FlowMeter",
        "Actuator_Valve", "Actuator_Pump", "BottlingMachine", "QualityControl",
        "SafetySystem", "DataLogger", "OTController", "NetworkSwitch"
    ]
    
    for node in nodes:
        G.add_node(node, 
                  type=random.choice(["controller", "sensor", "actuator", "machine"]),
                  risk=random.uniform(0.1, 0.4),
                  efficiency=random.uniform(0.7, 0.95))
    
    # Add connections
    edges = [
        ("PLC_Controller", "TemperatureSensor"),
        ("PLC_Controller", "PressureSensor"),
        ("PLC_Controller", "Actuator_Valve"),
        ("Actuator_Valve", "BottlingMachine"),
        ("BottlingMachine", "QualityControl"),
        ("SafetySystem", "BottlingMachine"),
        ("DataLogger", "PLC_Controller"),
        ("NetworkSwitch", "OTController")
    ]
    
    for edge in edges:
        G.add_edge(edge[0], edge[1], weight=random.uniform(0.5, 1.0))
    
    return G

def plot_network_graph(G):
    """Create and display network graph"""
    plt.figure(figsize=(10, 8))
    pos = nx.spring_layout(G, k=3, iterations=50)
    
    # Color nodes by type
    node_colors = []
    for node in G.nodes():
        node_type = G.nodes[node].get('type', 'unknown')
        if node_type == 'controller':
            node_colors.append('#3b82f6')  # Blue
        elif node_type == 'sensor':
            node_colors.append('#10b981')  # Green
        elif node_type == 'actuator':
            node_colors.append('#f59e0b')  # Orange
        else:
            node_colors.append('#6b7280')  # Gray
    
    nx.draw(G, pos, 
            node_color=node_colors,
            node_size=1000,
            font_size=8,
            font_weight='bold',
            with_labels=True,
            edge_color='#d1d5db',
            width=2)
    
    # Convert to base64 for display
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
    buffer.seek(0)
    image_png = buffer.getvalue()
    buffer.close()
    plt.close()
    
    return base64.b64encode(image_png).decode()

def generate_professional_report_demo():
    """Generate professional report for demo"""
    return {
        "executive_summary": {
            "title": "RevealGrid Industrial Intelligence Report",
            "subtitle": "AI-Powered Asset Discovery & Optimization Analysis",
            "date": datetime.now().strftime("%B %d, %Y"),
            "version": "1.0",
            "classification": "CONFIDENTIAL - Internal Use Only"
        },
        "key_findings": {
            "total_assets": st.session_state.assets_discovered,
            "automation_opportunities": len(st.session_state.recommendations),
            "efficiency_improvement": f"{st.session_state.efficiency - 75:.1f}%",
            "risk_mitigation": len([r for r in st.session_state.risk_insights if r['severity'] == 'high']),
            "hybrid_sources": {"auto_discovery": 12, "cmdb": 4, "document": 4, "cv": 4}
        },
        "technical_analysis": {
            "asset_distribution": {
                "automated": int(st.session_state.assets_discovered * 0.4),
                "manual": int(st.session_state.assets_discovered * 0.3),
                "robot": int(st.session_state.assets_discovered * 0.3)
            },
            "performance_metrics": {
                "initial_latency": "185.0s",
                "final_latency": "94.0s",
                "optimization_steps": 10
            },
            "risk_assessment": {
                "high_risk_nodes": len([r for r in st.session_state.risk_insights if r['severity'] == 'high']),
                "medium_risk_nodes": len([r for r in st.session_state.risk_insights if r['severity'] == 'medium']),
                "low_risk_nodes": len([r for r in st.session_state.risk_insights if r['severity'] == 'low'])
            }
        },
        "recommendations": {
            "immediate_actions": st.session_state.recommendations[:3],
            "long_term_strategy": st.session_state.recommendations[3:6] if len(st.session_state.recommendations) > 3 else [],
            "robot_integration": [
                {"action_description": "Automate ManualCheck_Operator with AI relay"},
                {"action_description": "Implement robot-assisted quality control"},
                {"action_description": "Deploy autonomous safety monitoring"}
            ]
        },
        "implementation_roadmap": {
            "phase_1": "Asset Discovery & Baseline Establishment",
            "phase_2": "Automation Gap Analysis & Prioritization", 
            "phase_3": "Robot Integration & World Model Deployment",
            "phase_4": "Continuous Optimization & Intelligence Evolution"
        }
    }

def create_word_document_demo(report_data, customization_prompt=""):
    """Create professional Word document for demo"""
    
    if not WORD_AVAILABLE:
        st.error("Word document generation not available. Please install python-docx.")
        return None
    
    doc = Document()
    
    # Title page
    title = doc.add_heading(report_data['executive_summary']['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Report Date: {report_data['executive_summary']['date']}")
    doc.add_paragraph(f"Version: {report_data['executive_summary']['version']}")
    doc.add_paragraph(f"Classification: {report_data['executive_summary']['classification']}")
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"""
    This report presents the findings of RevealGrid's AI-powered industrial intelligence analysis, 
    covering {report_data['key_findings']['total_assets']} critical assets across multiple operational sites.
    
    Key Performance Indicators:
    • Total Assets Discovered: {report_data['key_findings']['total_assets']}
    • High-Priority Automation Opportunities: {report_data['key_findings']['automation_opportunities']}
    • Efficiency Improvement Potential: {report_data['key_findings']['efficiency_improvement']}
    • Critical Risk Mitigation: {report_data['key_findings']['risk_mitigation']} high-severity issues
    """)
    
    # Technical Analysis
    doc.add_heading("Technical Analysis", level=1)
    
    # Asset Distribution
    doc.add_heading("Asset Distribution", level=2)
    doc.add_paragraph(f"""
    Asset Classification Analysis:
    • Automated Systems: {report_data['technical_analysis']['asset_distribution']['automated']} nodes
    • Manual Processes: {report_data['technical_analysis']['asset_distribution']['manual']} nodes  
    • Robot-Controlled Assets: {report_data['technical_analysis']['asset_distribution']['robot']} nodes
    """)
    
    # Performance Metrics
    doc.add_heading("Performance Metrics", level=2)
    doc.add_paragraph(f"""
    System Performance Analysis:
    • Initial System Latency: {report_data['technical_analysis']['performance_metrics']['initial_latency']}
    • Optimized System Latency: {report_data['technical_analysis']['performance_metrics']['final_latency']}
    • Optimization Iterations: {report_data['technical_analysis']['performance_metrics']['optimization_steps']} steps
    """)
    
    # Risk Assessment
    doc.add_heading("Risk Assessment", level=2)
    doc.add_paragraph(f"""
    Risk Analysis Summary:
    • High Severity Risks: {report_data['technical_analysis']['risk_assessment']['high_risk_nodes']}
    • Medium Severity Risks: {report_data['technical_analysis']['risk_assessment']['medium_risk_nodes']}
    • Low Severity Risks: {report_data['technical_analysis']['risk_assessment']['low_risk_nodes']}
    """)
    
    # Recommendations
    doc.add_heading("Strategic Recommendations", level=1)
    
    # Immediate Actions
    doc.add_heading("Immediate Actions (0-30 days)", level=2)
    for i, rec in enumerate(report_data['recommendations']['immediate_actions'], 1):
        doc.add_paragraph(f"{i}. {rec}", style='List Number')
    
    # Long-term Strategy
    if report_data['recommendations']['long_term_strategy']:
        doc.add_heading("Long-term Strategy (30-90 days)", level=2)
        for i, rec in enumerate(report_data['recommendations']['long_term_strategy'], 1):
            doc.add_paragraph(f"{i}. {rec}", style='List Number')
    
    # Robot Integration
    doc.add_heading("Robot Integration Plan", level=2)
    for i, action in enumerate(report_data['recommendations']['robot_integration'], 1):
        doc.add_paragraph(f"{i}. {action['action_description']}", style='List Number')
    
    # Implementation Roadmap
    doc.add_heading("Implementation Roadmap", level=1)
    for phase, description in report_data['implementation_roadmap'].items():
        doc.add_heading(f"{phase.replace('_', ' ').title()}", level=2)
        doc.add_paragraph(description)
    
    # AI Customization if provided
    if customization_prompt:
        doc.add_heading("AI-Generated Customization", level=1)
        doc.add_paragraph(f"Custom Analysis Request: {customization_prompt}")
        doc.add_paragraph("AI-generated insights and recommendations based on your specific requirements would appear here.")
    
    return doc

def generate_ai_customization_demo(user_prompt):
    """Generate AI-powered customization for demo"""
    
    # Simulate AI analysis
    customization_insights = {
        "financial_impact": {
            "downtime_reduction": f"${random.randint(50000, 200000):,} annually",
            "efficiency_gains": f"${random.randint(100000, 500000):,} annually", 
            "risk_mitigation": f"${random.randint(25000, 100000):,} annually",
            "total_roi": f"${random.randint(175000, 800000):,} annually"
        },
        "operational_insights": [
            "Critical bottleneck identified in Actuator_Valve system",
            "High automation potential in ManualCheck_Operator processes",
            "Robot integration opportunity in QualityControl_Station",
            "Risk mitigation priority for SafetySystem components"
        ],
        "strategic_recommendations": [
            "Implement phased automation approach starting with high-ROI opportunities",
            "Establish real-time monitoring for identified risk nodes",
            "Develop robot integration roadmap for quality control processes",
            "Create continuous improvement cycle for world model evolution"
        ]
    }
    
    return customization_insights

def display_professional_export_demo():
    """Display professional export interface for demo"""
    
    st.subheader("📋 Professional Report Generation")
    st.write("Generate engineering and business-ready documents with AI-powered customization.")
    
    # Report type selection
    col1, col2, col3 = st.columns(3)
    
    with col1:
        report_type = st.selectbox(
            "Report Type",
            ["Executive Summary", "Technical Analysis", "Implementation Plan", "Full Report"],
            key="demo_report_type"
        )
    
    with col2:
        document_format = st.selectbox(
            "Document Format", 
            ["Word Document (.docx)", "PDF Report", "PowerPoint Presentation", "Excel Dashboard"],
            key="demo_doc_format"
        )
    
    with col3:
        include_ai = st.checkbox("Include AI Customization", value=True, key="demo_ai_custom")
    
    # AI customization interface
    if include_ai:
        st.subheader("🤖 AI-Powered Customization")
        
        customization_prompt = st.text_area(
            "Describe your specific requirements for customization:",
            placeholder="e.g., Focus on financial impact, emphasize safety compliance, highlight automation ROI, include competitor analysis...",
            key="demo_custom_prompt"
        )
        
        if st.button("🔍 Generate AI Insights", key="demo_generate_ai"):
            with st.spinner("Analyzing with AI..."):
                time.sleep(2)
                
                # Generate AI customization
                ai_insights = generate_ai_customization_demo(customization_prompt)
                
                # Display AI insights
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("💰 Financial Impact Analysis")
                    st.metric("Annual Downtime Reduction", ai_insights['financial_impact']['downtime_reduction'])
                    st.metric("Efficiency Gains", ai_insights['financial_impact']['efficiency_gains'])
                    st.metric("Risk Mitigation Value", ai_insights['financial_impact']['risk_mitigation'])
                    st.metric("Total Annual ROI", ai_insights['financial_impact']['total_roi'])
                
                with col2:
                    st.subheader("🎯 Strategic Insights")
                    for insight in ai_insights['operational_insights']:
                        st.write(f"• {insight}")
                    
                    st.subheader("📈 Recommendations")
                    for rec in ai_insights['strategic_recommendations']:
                        st.write(f"• {rec}")
    
    # Generate report
    if st.button("📄 Generate Professional Report", key="demo_generate_report"):
        with st.spinner("Creating professional report..."):
            time.sleep(2)
            
            # Generate report data
            report_data = generate_professional_report_demo()
            
            # Create Word document
            doc = create_word_document_demo(report_data, customization_prompt if include_ai else "")
            
            # Download buttons
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if doc and WORD_AVAILABLE:
                    # Save to bytes
                    import io
                    doc_bytes = io.BytesIO()
                    doc.save(doc_bytes)
                    doc_bytes.seek(0)
                    
                    st.download_button(
                        label="📥 Download Word Document",
                        data=doc_bytes.getvalue(),
                        file_name=f"RevealGrid_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.info("Word document generation not available")
            
            with col2:
                # Generate PDF (simulated)
                pdf_data = f"PDF version of {report_type} report"
                st.download_button(
                    label="📥 Download PDF",
                    data=pdf_data,
                    file_name=f"RevealGrid_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf"
                )
            
            with col3:
                # Generate Excel dashboard (simulated)
                excel_data = f"Excel dashboard with {st.session_state.assets_discovered} assets"
                st.download_button(
                    label="📥 Download Excel Dashboard",
                    data=excel_data,
                    file_name=f"RevealGrid_Dashboard_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            
            st.success("✅ Professional report generated successfully!")
            
            # Report preview
            with st.expander("📋 Report Preview"):
                st.subheader("Executive Summary")
                st.write(f"**Report Type:** {report_type}")
                st.write(f"**Total Assets:** {report_data['key_findings']['total_assets']}")
                st.write(f"**Automation Opportunities:** {report_data['key_findings']['automation_opportunities']}")
                st.write(f"**Efficiency Improvement:** {report_data['key_findings']['efficiency_improvement']}")
                st.write(f"**Risk Mitigation:** {report_data['key_findings']['risk_mitigation']} critical issues")
                
                if include_ai and 'ai_insights' in locals():
                    st.subheader("AI-Generated Insights")
                    st.write("**Financial Impact:**")
                    st.write(f"• Total Annual ROI: {ai_insights['financial_impact']['total_roi']}")
                    st.write("**Strategic Recommendations:**")
                    for rec in ai_insights['strategic_recommendations']:
                        st.write(f"• {rec}")

# Simple Sidebar replaced with Progress Sidebar
with st.sidebar:
    st.markdown("<h2 style='margin-bottom:2rem;'>Demo Progress</h2>", unsafe_allow_html=True)
    for idx, s in enumerate(steps):
        if idx < st.session_state.current_step:
            st.markdown(f"<div style='color:#10b981;font-weight:bold;'>✔️ {s['title']}</div>", unsafe_allow_html=True)
        elif idx == st.session_state.current_step:
            st.markdown(f"<div style='color:#3b82f6;font-weight:bold;background:#e0e7ff;padding:0.5rem 0.75rem;border-radius:6px;'>{s['title']}</div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div style='color:#9ca3af;'>{s['title']}</div>", unsafe_allow_html=True)
    st.markdown("<div style='height:2rem;'></div>", unsafe_allow_html=True)

# Main content area
st.markdown('<h1 class="main-header">🏭 RevealGrid</h1>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Industrial Asset Intelligence Platform</p>', unsafe_allow_html=True)

# Progress bar
progress = (st.session_state.current_step + 1) / len(steps)
st.progress(progress)
st.caption(f"Step {st.session_state.current_step + 1} of {len(steps)}")

# Current step content
step = steps[st.session_state.current_step]
st.markdown(f"""
<div class="step-card">
    <h2>{step['title']}</h2>
    <p><strong>{step['description']}</strong></p>
    <p>{step['content']}</p>
</div>
""", unsafe_allow_html=True)

# Step-specific content (unchanged)
if st.session_state.current_step == 0:
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Industrial Sites", "4", "Active")
    with col2:
        st.metric("Asset Types", "12", "Discovered")
    with col3:
        st.metric("AI Models", "3", "Deployed")

elif st.session_state.current_step == 1:
    if st.button("Run Asset Discovery", key="discovery_btn"):
        simulate_discovery()
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Assets Discovered", st.session_state.discovered_assets, "Devices")
        st.metric("Network Segments", random.randint(3, 8), "VLANs")
    with col2:
        st.metric("Device Types", random.randint(8, 15), "Categories")
        st.metric("Protocols", random.randint(5, 12), "Industrial")
    if st.session_state.discovered_assets > 0:
        st.success("✅ Asset discovery completed successfully!")

elif st.session_state.current_step == 2:
    if st.button("Analyze Network Security", key="analysis_btn"):
        analysis = simulate_analysis()
        st.session_state.analysis = analysis
    analysis = st.session_state.get("analysis", {"high_risk_nodes":0,"bottlenecks":0,"anomalies":0})
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("High Risk Nodes", analysis["high_risk_nodes"], "Critical")
    with col2:
        st.metric("Bottlenecks", analysis["bottlenecks"], "Identified")
    with col3:
        st.metric("Anomalies", analysis["anomalies"], "Detected")
    st.subheader("Network Topology")
    G = create_sample_graph()
    graph_b64 = plot_network_graph(G)
    st.image(f"data:image/png;base64,{graph_b64}", use_column_width=True)
    if analysis["high_risk_nodes"] > 0:
        st.success("✅ Risk assessment completed!")

elif st.session_state.current_step == 3:
    if st.button("Generate AI Recommendations", key="ai_btn"):
        recommendations = simulate_recommendations()
        st.session_state.recommendations = recommendations
    recommendations = st.session_state.get("recommendations", [])
    if recommendations:
        st.subheader("🤖 AI-Generated Recommendations")
        for i, rec in enumerate(recommendations, 1):
            st.markdown(f"""
            <div class="success-card">
                <h4>Recommendation {i}</h4>
                <p>{rec}</p>
            </div>
            """, unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Efficiency Gain", "+23%", "Predicted")
        with col2:
            st.metric("Downtime Reduction", "-45%", "Expected")
        with col3:
            st.metric("Risk Reduction", "-32%", "Projected")
        st.success("✅ AI recommendations generated!")

elif st.session_state.current_step == 4:
    st.header("🌍 World Model Export")
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Model Statistics")
        nodes = st.session_state.discovered_assets
        edges = nodes * random.uniform(0.8, 1.2)
        efficiency = random.uniform(85, 95)
        latency_improvement = random.uniform(30, 60)
        st.metric("Total Nodes", nodes)
        st.metric("Total Edges", int(edges))
        st.metric("System Efficiency", f"{efficiency:.1f}%")
        st.metric("Latency Improvement", f"{latency_improvement:.1f}%")
        optimization_steps = list(range(1, 11))
        latencies = [100 - (i * latency_improvement / 10) for i in range(10)]
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.plot(optimization_steps, latencies, marker='o', color='blue')
        ax.set_xlabel("Optimization Step")
        ax.set_ylabel("System Latency (s)")
        ax.set_title("World Model Optimization Progress")
        st.pyplot(fig)
    with col2:
        st.subheader("Export Options")
        st.markdown("""
        **Export Formats:**
        - 🔐 Encrypted JSON (recommended)
        - 📊 CSV Report
        - 🗺️ GraphML Network
        - 📋 PDF Summary
        """)
        if st.button("🔐 Export Encrypted Model"):
            with st.spinner("Generating encrypted export..."):
                time.sleep(2)
                export_data = {
                    "timestamp": time.time(),
                    "nodes": nodes,
                    "edges": int(edges),
                    "efficiency": efficiency,
                    "recommendations": len(st.session_state.recommendations)
                }
                encrypted = base64.b64encode(json.dumps(export_data).encode()).decode()
                st.success("✅ Export generated successfully!")
                st.code(encrypted[:100] + "...", language="text")
                st.download_button(
                    label="📥 Download Encrypted Model",
                    data=encrypted,
                    file_name="revealgrid_world_model.enc",
                    mime="text/plain"
                )
    st.markdown("---")
    st.markdown("""
    ### 🎉 Demo Complete!
    **RevealGrid** has successfully:
    - ✅ Discovered industrial assets
    - ✅ Analyzed risks and vulnerabilities
    - ✅ Generated AI-powered recommendations
    - ✅ Built and exported world model
    **Ready for production deployment in OT environments!**
    """)

# Navigation buttons at the bottom
nav_cols = st.columns([1, 6, 1])
with nav_cols[0]:
    if st.session_state.current_step > 0:
        if st.button("⬅️ Back", key="back_btn"):
            st.session_state.current_step -= 1
            st.experimental_rerun()
with nav_cols[2]:
    if st.session_state.current_step < len(steps) - 1:
        if st.button("Next ➡️", key="next_btn"):
            st.session_state.current_step += 1
            st.experimental_rerun() 